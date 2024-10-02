import os
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from docx import Document

# Configura o caminho do Tesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Caminho do PDF
pdf_path = r'C:\Users\Dell\Desktop\teste_imagem_para_texto\teste\convencao_removed.pdf'

# Diretório temporário para salvar as imagens das páginas do PDF
temp_dir = r'C:\Users\Dell\Desktop\teste_imagem_para_texto\teste\temp_images'

# Cria o diretório temporário, se não existir
if not os.path.exists(temp_dir):
    os.makedirs(temp_dir)

# Converter as páginas do PDF em imagens
paginas_imagens = convert_from_path(pdf_path, 300, output_folder=temp_dir)

# Cria um documento Word
doc = Document()

# Itera pelas imagens das páginas do PDF
for i, imagem in enumerate(paginas_imagens):
    # Salva cada página como imagem (opcional se você quiser ver as imagens salvas)
    
    imagem_path = os.path.join(temp_dir, f'pagina_{i+1}.png')
    imagem.save(imagem_path, 'PNG')

    # Realiza OCR na imagem para extrair o texto
    texto = pytesseract.image_to_string(imagem)
    # Garantindo que o texto esteja em UTF-8
    texto = texto.encode('utf-8', errors='ignore').decode('utf-8')

    # Adiciona o texto extraído ao documento
    doc.add_paragraph(f"Texto extraído da página {i+1}:")
    doc.add_paragraph(texto)
    doc.add_paragraph("\n\n")

# Salva o documento .docx
doc.save(os.path.join(temp_dir, 'texto_extraido_do_pdf.docx'))

# Opcional: Limpar o diretório temporário (remover as imagens)
for arquivo in os.listdir(temp_dir):
    if arquivo.endswith('.ppm') or arquivo.endswith('.png'):
        os.remove(os.path.join(temp_dir, arquivo))

print("Processo concluído! O texto foi salvo em 'texto_extraido_do_pdf.docx'.")
