import os
from PIL import Image
import pytesseract
from docx import Document

# Caminho do Tesseract
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Caminho da pasta contendo as imagens
pasta_imagens = r'C:\Users\Dell\Desktop\teste_imagem_para_texto\teste'

# Cria um documento Word
doc = Document()

# Itera pelos arquivos da pasta
for arquivo in os.listdir(pasta_imagens):
    # Verifica se o arquivo é uma imagem (extensões comuns)
    if arquivo.endswith(('.png', '.jpg', '.jpeg', '.bmp', '.tiff')):
        img_path = os.path.join(pasta_imagens, arquivo)
        print(f'Lendo arquivo: {img_path}')

        # Abre a imagem
        img = Image.open(img_path)

        # Extrai o texto da imagem
        texto = pytesseract.image_to_string(img)

        # Adiciona o texto ao documento
        doc.add_paragraph(f"Texto extraído da imagem: {arquivo}")
        doc.add_paragraph(texto)
        doc.add_paragraph("\n\n")  # Adiciona um espaço entre os textos das imagens

# Salva o documento .docx
doc.save(os.path.join(pasta_imagens, 'texto_extraido_imagens.docx'))

print("Processo concluído! O texto foi salvo em 'texto_extraido_imagens.docx'.")
